import os
import pandas as pd
from docx import Document
import unicodedata
import re
import zipfile
import json
from datetime import datetime

def substituir_texto(container, dicionario_dados):
    """Substitui texto em parágrafos e tabelas mantendo formatação."""
    for p in container.paragraphs:
        for busca, substituto in dicionario_dados.items():
            if busca in p.text:
                for run in p.runs:
                    if busca in run.text:
                        run.text = run.text.replace(busca, str(substituto))
    
    if hasattr(container, 'tables'):
        for tabela in container.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    substituir_texto(celula, dicionario_dados)

def limpar_nome_pasta(nome):
    """Normaliza nome para criar pastas seguras (sem acento, minúsculo)."""
    nfkd = unicodedata.normalize('NFKD', nome)
    nome_sem_acento = "".join([c for c in nfkd if not unicodedata.combining(c)])
    nome_limpo = re.sub(r'[^a-zA-Z0-9 ]', '', nome_sem_acento)
    return nome_limpo.lower().replace(' ', '_')

def criar_zip(caminho_pasta, caminho_zip):
    """Compacta todo o conteúdo da pasta em um arquivo .zip"""
    with zipfile.ZipFile(caminho_zip, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for root, dirs, files in os.walk(caminho_pasta):
            for file in files:
                caminho_completo = os.path.join(root, file)
                # Adiciona o arquivo ao zip mantendo a estrutura relativa
                zipf.write(caminho_completo, os.path.relpath(caminho_completo, os.path.join(caminho_pasta, '..')))

def processar_lote():
    arquivo_excel = "entradas/respostas_forms.xlsx"
    
    if not os.path.exists(arquivo_excel):
        print("ERRO: Excel não encontrado.")
        return

    print("--- Iniciando Processamento ---")
    df = pd.read_excel(arquivo_excel)

    for index, linha in df.iterrows():
        nome = linha['Nome Completo']
        cpf = linha['CPF']
        curso = linha['Curso']
        data = linha['Data de Conclusão']
        
        print(f"Processando: {nome}...")

        # Estrutura do LOG individual
        log_execucao = {
            "colaborador": nome,
            "data_processamento": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "arquivos_gerados": [],
            "erros": []
        }

        # Preparação de Pastas
        nome_pasta = limpar_nome_pasta(nome)
        caminho_pasta_saida = os.path.join("saidas", nome_pasta)
        if not os.path.exists(caminho_pasta_saida):
            os.makedirs(caminho_pasta_saida)

        # Geração do Certificado
        dados_certificado = {
            "{{NOME}}": nome, "{{CPF}}": cpf,
            "{{CURSO}}": curso, "{{DATA}}": data
        }
        
        caminho_template = f"templates/{curso}.docx"
        # Fallback para teste se não tiver o template específico
        if not os.path.exists(caminho_template):
            caminho_template = "templates/NR06.docx" 

        if os.path.exists(caminho_template):
            try:
                doc = Document(caminho_template)
                substituir_texto(doc, dados_certificado)
                nome_arquivo_docx = f"{nome_pasta}_{curso}.docx"
                caminho_final_docx = os.path.join(caminho_pasta_saida, nome_arquivo_docx)
                doc.save(caminho_final_docx)
                
                log_execucao["arquivos_gerados"].append(nome_arquivo_docx)
            except Exception as e:
                log_execucao["erros"].append(str(e))
        else:
            log_execucao["erros"].append(f"Template {curso} não encontrado")

        # Salva o Relatório JSON na pasta
        with open(os.path.join(caminho_pasta_saida, 'relatorio.json'), 'w', encoding='utf-8') as f:
            json.dump(log_execucao, f, indent=4, ensure_ascii=False)

        # Cria o ZIP final
        caminho_zip = os.path.join("saidas", f"{nome_pasta}.zip")
        criar_zip(caminho_pasta_saida, caminho_zip)
        print(f"   -> ZIP criado: {nome_pasta}.zip")

if __name__ == "__main__":
    processar_lote()